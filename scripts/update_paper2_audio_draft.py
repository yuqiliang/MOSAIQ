#!/usr/bin/env python3
"""Create a versioned Paper 2 draft with the ISD audio reference extension."""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


def paragraph_starting(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def heading(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Heading not found: {text}")


def insert_after(paragraph: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    inserted.style = style
    inserted.add_run(text)
    return inserted


def apply_paragraph_template(target: Paragraph, source: Paragraph) -> None:
    """Copy the reference paragraph and first-run formatting."""
    if target._p.pPr is not None:
        target._p.remove(target._p.pPr)
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))
    if target.runs and source.runs and source.runs[0]._r.rPr is not None:
        target.runs[0]._r.insert(0, deepcopy(source.runs[0]._r.rPr))


def insert_heading_after(
    paragraph: Paragraph,
    text: str,
    template: Paragraph,
) -> Paragraph:
    inserted = insert_after(paragraph, text, "Heading 2")
    apply_paragraph_template(inserted, template)
    return inserted


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    heading_two_template = heading(document, "Evaluation and robustness procedures")

    abstract = heading(document, "Abstract")._p.getnext()
    abstract_paragraph = Paragraph(abstract, heading(document, "Abstract")._parent)
    abstract_paragraph.text = (
        abstract_paragraph.text
        + " A separately versioned ISD audio extension verifies seven source "
        "archives and freezes 820 waveform-linked clips; clip- and response-level "
        "Target Mean and descriptor Ridge references add 2,988 held-out "
        "predictions with clip-clustered uncertainty without changing the "
        "no-audio v0.1 paper freeze."
    )

    scope = paragraph_starting(document, "MOSAIQ v0.1 contributes")
    scope.text = (
        "MOSAIQ v0.1 contributes a versioned benchmark contract rather than a "
        "new prediction architecture. It provides seven task definitions, "
        "deterministic and leakage-aware split version 0.1.0, eleven "
        "task/dataset manifests with record hashes, a 49-check validation "
        "report, seventeen reproducible tabular reference experiments, model "
        "and data cards, and robustness analyses. The fixed v0.1.0-dev release "
        "and its Paper 2 evidence remain no-audio, and those runs record "
        "audio_used=false. A separate ISD 0.1.0-audio extension now tests "
        "waveform acquisition, QC, relative descriptors, reference prediction, "
        "and uncertainty on the same frozen clip partitions. It is not a "
        "four-dataset audio or multimodal release."
    )

    interface = paragraph_starting(document, "The unified interface resolves")
    interface.text = interface.text.replace(
        "Every run records software versions and audio_used=false.",
        "Every tabular run records software versions and audio_used=false; "
        "the separately versioned ISD audio runs record audio_used=true.",
    )

    evaluation = paragraph_starting(document, "Test uncertainty is estimated")
    inserted = insert_heading_after(
        evaluation,
        "ISD audio reference extension",
        heading_two_template,
    )
    inserted = insert_after(
        inserted,
        "ISD waveforms were acquired from Zenodo record 10672568, source "
        "version 1.0.1-alpha.1, under the record's CC BY 4.0 declaration. Seven "
        "benchmark-candidate archives were verified against the source MD5 "
        "values and extracted outside Git. Candidate WAV members were linked "
        "to MOSAIQ clips using normalised GroupID and source-location evidence. "
        "Ambiguous, missing, unmatched, duplicate, and frozen-excluded records "
        "were retained in a versioned exclusion table rather than imputed.",
        "Normal",
    )
    insert_after(
        inserted,
        "Technical QC checks readability, finite and non-zero samples, channel "
        "count, sample rate, duration agreement, encoding, and amplitude-scale "
        "warnings. Deterministic waveform descriptors include crest factor, "
        "zero-crossing rate, stereo correlation, spectral centroid, bandwidth, "
        "85% rolloff, and normalised low-, mid-, and high-band power. RMS and "
        "peak are retained for QC but excluded from Ridge pending confirmation "
        "of calibrated sound-pressure metadata. Target Mean and multi-output "
        "Ridge are fitted on train only for clip-mean and individual-response "
        "ISO coordinates. Response records inherit clip partitions, and all "
        "uncertainty resamples are clustered by clip_id.",
        "Normal",
    )

    repository = paragraph_starting(document, "The repository root contains")
    repository.text = repository.text.replace(
        "governance records, and the result submission contract.",
        "governance records, and the result submission contract. "
        "benchmark/audio/ separately contains the ISD source registry, audio "
        "manifest, QC, frozen cohort and exclusions, descriptors, predictions, "
        "uncertainty outputs, model cards, and audio validator.",
    )

    generated = paragraph_starting(document, "The baseline result table has")
    inserted = insert_heading_after(
        generated,
        "ISD audio records",
        heading_two_template,
    )
    insert_after(
        inserted,
        "The audio manifest contains 1,021 source or expected-clip rows. The "
        "frozen 0.1.0-audio cohort accepts 820 unique waveform-to-clip mappings "
        "(546 train, 154 development, and 120 test) and records 201 exclusions: "
        "143 missing source assets, 52 unmatched source assets, three ambiguous "
        "mappings, one byte-identical duplicate, and two frozen-split "
        "exclusions. Source archive MD5 values, extracted WAV SHA-256 values, "
        "cohort checksums, and prediction-level identifiers provide an auditable "
        "chain from Zenodo file to held-out result.",
        "Normal",
    )

    media_warning = paragraph_starting(document, "The warnings are retained")
    media_warning.text = media_warning.text.replace(
        "Asset identifiers are complete in the tabular resources, but media "
        "files are not materialised. Consequently, structural asset coverage "
        "must not be interpreted as waveform or video availability in the "
        "repository.",
        "Asset identifiers are complete in the tabular resources, while raw "
        "media are not part of the no-audio repository package. A separately "
        "versioned ISD working track materialises rights-cleared waveforms "
        "outside Git; it must not be interpreted as audio availability for "
        "ARAUS, SATP, DeLTA, or the visual modality.",
    )

    last_validation = paragraph_starting(
        document,
        "Figure 7. Direction-normalised paired RMSE improvement",
    )
    inserted = insert_heading_after(
        last_validation,
        "ISD audio technical validation",
        heading_two_template,
    )
    inserted.paragraph_format.page_break_before = True
    inserted = insert_after(
        inserted,
        "All 820 accepted WAV files are readable, finite, non-empty, non-silent, "
        "stereo, and within 0.005 s of the harmonised clip duration. The cohort "
        "contains 341 files at 44.1 kHz and 479 at 48 kHz. Actual encoding is "
        "float32 for 769 files and int16 for 51 Groningen files. No waveform "
        "SHA-256 crosses train, development, and test. All files retain a "
        "calibration warning; 247 float files require amplitude-scale review, "
        "and the Groningen encoding differs from the source prose.",
        "Normal",
    )
    insert_after(
        inserted,
        "At clip level, descriptor Ridge improves test Eventfulness RMSE from "
        "0.3198 to 0.2924 (paired improvement 0.0274; 95% clip-clustered CI "
        "0.0138 to 0.0421), but does not improve test Pleasantness (0.2812 to "
        "0.3043; improvement -0.0231; 95% CI -0.0485 to 0.0033). At response "
        "level, Ridge improves test Eventfulness RMSE by 0.0271 (95% CI 0.0137 "
        "to 0.0399), whereas the Pleasantness difference is inconclusive. These "
        "mixed results validate the audio interface and motivate learned "
        "baselines; they are not a state-of-the-art performance claim.",
        "Normal",
    )

    current_runs = paragraph_starting(document, "•  All current runs are tabular")
    current_runs.text = (
        "•  The fixed v0.1 Paper 2 runs are tabular and record audio_used=false; "
        "the separate ISD audio extension records audio_used=true."
    )

    media = paragraph_starting(document, "Asset columns identify source media")
    media.text = (
        "The no-audio v0.1 candidate does not include waveform or video "
        "payloads. The separate ISD 0.1.0-audio extension makes a rights-cleared "
        "820-clip working cohort available outside Git and executes only Target "
        "Mean and descriptor Ridge references. Learned log-mel CNN, pretrained "
        "encoder, visual-only, fusion, missing-modality, and added-noise "
        "evaluations remain future tasks. A media-bearing public release must "
        "increment the benchmark version, complete the rights audit, publish "
        "media checksums, define feature-extraction versions, and rerun the full "
        "validation and manuscript output pipeline. Results must not be "
        "retroactively attached to v0.1.0-dev."
    )

    availability = paragraph_starting(document, "The development repository is")
    availability.text = (
        availability.text
        + " The ISD audio working track cites Zenodo record "
        "https://doi.org/10.5281/zenodo.10672568 and remains separate from the "
        "prepared no-audio RDR archive; no UCL RDR audio record or upload has "
        "been created."
    )
    for heading_text in ("Data Availability", "Code Availability", "Acknowledgements"):
        heading_paragraph = heading(document, heading_text)
        heading_paragraph.paragraph_format.page_break_before = False
        spacer = insert_after(heading_paragraph, "", "Normal")
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = Pt(1)

    references = heading(document, "References")
    reference_started = False
    for paragraph in document.paragraphs:
        if paragraph._p is references._p:
            reference_started = True
            continue
        if not reference_started or not paragraph.text.strip():
            continue
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.size = Pt(9)

    paragraph_starting(
        document,
        "Table 2. Materialised MOSAIQ v0.1 resources.",
    ).paragraph_format.page_break_before = True

    document.core_properties.title = (
        "MOSAIQ Paper 2 Scientific Data draft v0.2.0 - ISD audio reference"
    )
    document.core_properties.subject = (
        "Versioned manuscript draft; no-audio fixed outputs plus separate ISD "
        "audio reference extension"
    )
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(f"WROTE {args.output}")


if __name__ == "__main__":
    main()
